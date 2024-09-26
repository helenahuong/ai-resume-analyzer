import io
import re
import spacy
from PyPDF2 import PdfReader
from docx import Document
from fastapi import UploadFile

nlp = spacy.load("en_core_web_sm")

class ResumeParser:
    def __init__(self, data, is_file=True):
        self.data = data
        self.is_file = is_file

    async def parse(self):
        content = await self.data.read()
        text = self.extract_text(self.data.filename, content)
        return self.extract_data(text)

    async def parse_text(self):
        text = self.data
        return self.extract_data(text)

    def extract_text(self, filename, content):
        if filename.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(content)
        elif filename.lower().endswith(('.doc', '.docx')):
            return self.extract_text_from_docx(content)
        else:
            return content.decode('utf-8', errors='ignore')

    def extract_text_from_pdf(self, content):
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def extract_text_from_docx(self, content):
        document = Document(io.BytesIO(content))
        return "\n".join([para.text for para in document.paragraphs])

    def extract_data(self, text):
        doc = nlp(text)
        keywords = [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and token.is_alpha]
        email = self.extract_email(text)
        phone = self.extract_phone_number(text)
        name = self.extract_name(doc)
        return {
            "name": name,
            "email": email,
            "phone": phone,
            "keywords": keywords,
        }

    def extract_email(self, text):
        match = re.search(r'[\w\.-]+@[\w\.-]+', text)
        return match.group(0) if match else None

    def extract_phone_number(self, text):
        match = re.search(r'\+?\d[\d -]{8,}\d', text)
        return match.group(0) if match else None

    def extract_name(self, doc):
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                return ent.text
        return None
# backend/app/services/resume_parser.py

import io
import re
import spacy
from PyPDF2 import PdfReader
from docx import Document
from fastapi import UploadFile

nlp = spacy.load("en_core_web_sm")

class ResumeParser:
    def __init__(self, file: UploadFile):
        self.file = file

    async def parse(self):
        content = await self.file.read()
        text = self.extract_text(self.file.filename, content)
        data = self.extract_data(text)
        return data

    def extract_text(self, filename, content):
        if filename.lower().endswith('.pdf'):
            return self.extract_text_from_pdf(content)
        elif filename.lower().endswith(('.doc', '.docx')):
            return self.extract_text_from_docx(content)
        else:
            return content.decode('utf-8', errors='ignore')

    def extract_text_from_pdf(self, content):
        reader = PdfReader(io.BytesIO(content))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text

    def extract_text_from_docx(self, content):
        document = Document(io.BytesIO(content))
        return "\n".join([para.text for para in document.paragraphs])

    def extract_data(self, text):
        doc = nlp(text)
        keywords = [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and token.is_alpha]
        email = self.extract_email(text)
        phone = self.extract_phone_number(text)
        name = self.extract_name(doc)
        return {
            "name": name,
            "email": email,
            "phone": phone,
            "keywords": keywords,
            "text": text,
        }

    def extract_email(self, text):
        match = re.search(r'[\w\.-]+@[\w\.-]+', text)
        return match.group(0) if match else None

    def extract_phone_number(self, text):
        match = re.search(r'\+?\d[\d -]{8,}\d', text)
        return match.group(0) if match else None

    def extract_name(self, doc):
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                return ent.text
        return None
